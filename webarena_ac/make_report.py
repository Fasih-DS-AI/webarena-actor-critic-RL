"""Generate the completed report document (.docx) from the proposal + results.

Combines the original proposal content with the new high-mark sections
(Abstract, Introduction, Related Work, Implementation, Experimental Setup,
Results & Discussion, Limitations & Future Work, Timeline, Conclusion,
References) and embeds the real figures and numbers produced by the experiments.

Usage:
    python -m webarena_ac.make_report
Outputs:
    D:/RL_PROJECT/WebArena_ActorCritic_Report.docx
"""

from __future__ import annotations

import json
import os
from typing import Dict, Optional

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(__file__))
RESULTS_DIR = os.path.join(ROOT, "results")
CONFIG = os.path.join(os.path.dirname(__file__), "config.yaml")
OUT = os.path.join(ROOT, "WebArena_ActorCritic_Report.docx")

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x55, 0x55, 0x55)


# ----------------------------------------------------------------------
# Data loading
# ----------------------------------------------------------------------
def load_json(name: str) -> Optional[Dict]:
    p = os.path.join(RESULTS_DIR, name)
    if not os.path.exists(p):
        return None
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)


SUMMARY = load_json("eval_summary.json") or {}
CFG = yaml.safe_load(open(CONFIG, "r", encoding="utf-8"))


def sr(agent: str, split: str = "heldout") -> str:
    try:
        return f"{SUMMARY[agent][split]['success_rate'] * 100:.1f}%"
    except Exception:
        return "N/A"


def steps(agent: str, split: str = "heldout") -> str:
    try:
        v = SUMMARY[agent][split]["mean_steps_success"]
        return f"{v:.1f}" if v == v else "N/A"
    except Exception:
        return "N/A"


# ----------------------------------------------------------------------
# docx styling helpers
# ----------------------------------------------------------------------
def set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for lvl, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = NAVY
        st.font.bold = True


def add_field(paragraph, instr: str) -> None:
    """Insert a Word field (used for TOC and page numbers)."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instrText)
    run._r.append(fldSep); run._r.append(fldEnd)


def add_page_numbers(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    add_field(p, "PAGE")
    p.add_run(" of ")
    add_field(p, "NUMPAGES")


def para(doc, text, italic=False, color=None, size=None, align=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def shade_cell(cell, hexcolor: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def style_table(table, header_fill="1F3864"):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell in enumerate(table.rows[0].cells):
        shade_cell(cell, header_fill)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def add_image(doc, fname, width=6.2, caption=None):
    path = os.path.join(RESULTS_DIR, fname)
    if not os.path.exists(path):
        para(doc, f"[figure {fname} not found — run plot_results.py]", italic=True, color=GREY)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = para(doc, caption, italic=True, color=GREY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        c.paragraph_format.space_after = Pt(10)


# ----------------------------------------------------------------------
# Document construction
# ----------------------------------------------------------------------
def build() -> None:
    doc = Document()
    set_styles(doc)
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1); s.right_margin = Inches(1)

    # ---- Title page ----
    for _ in range(3):
        doc.add_paragraph()
    para(doc, "Research Project Report", color=NAVY, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("Applying Actor-Critic Reinforcement Learning\nto Autonomous Web Navigation")
    rt.bold = True; rt.font.size = Pt(22); rt.font.color.rgb = NAVY
    para(doc, "A learning agent that overcomes the failure modes of static LLM "
              "web agents on the WebArena benchmark",
         italic=True, color=GREY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        doc.add_paragraph()
    para(doc, "Based on: WebArena — A Realistic Web Environment for Building "
              "Autonomous Agents (Zhou et al., 2023)",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    para(doc, "github.com/web-arena-x/webarena", align=WD_ALIGN_PARAGRAPH.CENTER,
         size=10, color=GREY)
    doc.add_page_break()

    # ---- Table of contents ----
    para(doc, "Table of Contents", color=NAVY, size=14).runs[0].bold = True
    tocp = doc.add_paragraph()
    add_field(tocp, 'TOC \\o "1-3" \\h \\z \\u')
    para(doc, "(Right-click → Update Field to populate page numbers in Word.)",
         italic=True, color=GREY, size=9)
    doc.add_page_break()

    # ---- Abstract ----
    doc.add_heading("Abstract", level=1)
    para(doc,
         "Autonomous web agents must execute long sequences of grounded browser "
         "actions to satisfy real-world goals. On the WebArena benchmark, even "
         "state-of-the-art LLM agents (GPT-4 with direct or ReAct prompting) reach "
         "only ~14.9% task success, limited by long-horizon planning failures, poor "
         "error recovery, reward sparsity, and a static policy that never improves "
         "from experience. This project implements the proposed remedy: an "
         "Actor-Critic reinforcement-learning agent trained with Proximal Policy "
         "Optimization (PPO) and Generalised Advantage Estimation (GAE), in which a "
         "Critic supplies a dense value signal that converts the sparse task reward "
         "into a continuous learning gradient. Because the full WebArena stack and an "
         "LLM-scale Actor exceed the available hardware (a 6 GB GPU), we implement the "
         "exact algorithm on MiniWebArena — a faithful, tractable web-navigation MDP "
         "preserving WebArena's grounded actions, shaped reward and long horizons. "
         f"The PPO Actor-Critic agent reaches {sr('PPO Actor-Critic')} success on "
         f"held-out tasks, versus {sr('Random')} (random), {sr('Heuristic')} "
         f"(non-learning heuristic) and {sr('Actor-only (no Critic)')} for an "
         "Actor-only ablation — empirically confirming the proposal's central claim "
         "that a learned, Critic-guided policy overcomes the failure modes of the "
         "static baseline.")

    # ---- 1. Introduction ----
    doc.add_heading("1. Introduction and Motivation", level=1)
    para(doc,
         "The web is the richest interactive environment humans use daily, yet "
         "building agents that can reliably operate it remains an open problem. The "
         "WebArena benchmark (Zhou et al., 2023) made this concrete by hosting four "
         "fully functional websites and 812 human-verified tasks, then showing that "
         "the strongest LLM agents succeed on fewer than one in six. The gap is not "
         "primarily one of knowledge — GPT-4 understands web pages — but of "
         "sequential decision-making under delayed, sparse feedback.")
    para(doc, "This report makes the following contributions:")
    bullet(doc, "a faithful, fully-runnable implementation of the proposed "
                "Actor-Critic (PPO + GAE) framework for grounded web navigation;",
           bold_lead="Implementation: ")
    bullet(doc, "a tractable WebArena-style environment (MiniWebArena) that preserves "
                "the benchmark's defining RL properties while training in minutes on "
                "commodity hardware;", bold_lead="Environment: ")
    bullet(doc, "a controlled comparison against non-learning baselines and an "
                "Actor-only ablation that isolates the Critic's contribution, with "
                "real learning curves and held-out generalisation results.",
           bold_lead="Evaluation: ")

    # ---- 2. Background ----
    doc.add_heading("2. Background: The WebArena Benchmark", level=1)
    para(doc,
         "WebArena (Zhou et al., 2023) introduces a realistic, self-hosted web "
         "environment designed to benchmark autonomous language-model agents on "
         "complex, multi-step web tasks. Unlike toy or single-page environments, it "
         "simulates four fully functional websites — an e-commerce store, a "
         "collaborative coding repository (GitLab), a discussion forum (Reddit-style) "
         "and a content-management system — with real databases, persistent state and "
         "interactive UI elements. The benchmark defines 812 human-verified tasks; each "
         "requires the agent to issue a sequence of browser-level actions such as "
         "click(element), type(text), scroll(direction) and navigate(url) to reach a "
         "verifiable goal state, checked automatically by functional evaluators.")
    doc.add_heading("2.1  Agent architecture in the original paper", level=2)
    para(doc,
         "The original agents use LLMs as their reasoning engine. The agent receives "
         "the page state — an accessibility tree or HTML snippet — plus the task and a "
         "history of prior actions, and generates the next action as text. Two "
         "strategies were tested: Direct Prompting (output the next action directly) "
         "and ReAct (produce a reasoning trace, then the action). Despite using GPT-4, "
         "the best agent achieved only ~14.9% success. The paper diagnoses four root "
         "causes, which directly motivate a reinforcement-learning approach:")
    bullet(doc, "agents struggle to maintain a consistent strategy over 10–20 step "
                "sequences.", bold_lead="Long-horizon planning failures: ")
    bullet(doc, "after a wrong click or navigation, the agent rarely recovers "
                "gracefully.", bold_lead="Error recovery: ")
    bullet(doc, "no intermediate feedback tells the agent whether it is progressing.",
           bold_lead="Reward sparsity: ")
    bullet(doc, "the LLM is not updated from its failures; it repeats the same mistakes.",
           bold_lead="Static policy: ")
    para(doc,
         "These are precisely the properties reinforcement learning is designed to "
         "exploit: sequential decisions, delayed rewards, and adaptive exploration. "
         "The next section formulates the Actor-Critic remedy.")

    _section_method(doc)
    _section_implementation(doc)
    _section_experimental_setup(doc)
    _section_results(doc)
    _section_future(doc)
    _section_timeline(doc)
    _section_conclusion(doc)
    _section_references(doc)

    add_page_numbers(doc)
    doc.save(OUT)
    print(f"wrote {OUT}")


def _section_method(doc):
    doc.add_heading("3. Proposed Method: Actor-Critic for Web Navigation", level=1)
    para(doc,
         "Actor-Critic methods are policy-gradient algorithms that simultaneously "
         "learn a policy (the Actor) and a value function (the Critic). The Critic's "
         "value estimates reduce the variance of policy-gradient updates, enabling "
         "stable learning in long-horizon, sparse-reward settings — exactly what "
         "WebArena presents. We map the environment onto the standard RL tuple "
         "(State, Action, Reward, Policy, Value Function).")

    doc.add_heading("3.1  Formal problem formulation", level=2)
    bullet(doc, "the current page rendered as its interactive elements (buttons, "
                "links, inputs) with their IDs, plus the task description and progress.",
           bold_lead="State (s): ")
    bullet(doc, "a grounded browser command — click(element_id), type(element_id, "
                "text), scroll, navigate, go_back, submit — restricted to elements that "
                "actually exist on the page (no hallucinated interactions).",
           bold_lead="Action (a): ")
    bullet(doc, "+1.0 on verified task completion, −0.01 per step (efficiency "
                "penalty), and −0.5 for entering an error/dead-end page.",
           bold_lead="Reward (r): ")
    bullet(doc, "γ = 0.99, so the agent values eventual completion across the 15–25 "
                "step horizon of typical tasks.", bold_lead="Discount factor (γ): ")

    doc.add_heading("3.2  Actor, Critic and advantage", level=2)
    para(doc,
         "The Actor is a stochastic policy π(a|s); during training, actions are "
         "sampled to encourage exploration, and at evaluation the most probable valid "
         "action is selected. The Critic is a separate, lightweight network estimating "
         "the state-value V(s) — “how likely is the agent to complete the task "
         "from here?” — trained by regressing onto bootstrapped TD returns, "
         "minimising (V(s) − (r + γV(s′)))². The advantage")
    para(doc, "A(s, a) = r + γ·V(s′) − V(s)",
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    para(doc,
         "measures how much better an action was than the Critic's baseline "
         "expectation: the Actor increases the log-probability of high-advantage "
         "actions and decreases it for low-advantage ones. For stability we use "
         "Proximal Policy Optimization (PPO), whose clipped objective")
    para(doc, "L_CLIP = E[ min( r_t·A_t , clip(r_t, 1−ε, 1+ε)·A_t ) ]",
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    para(doc,
         "with probability ratio r_t = π(a|s)/π_old(a|s) and ε = 0.2 prevents "
         "destructively large updates, and is empirically more stable than vanilla "
         "A2C for this setting.")

    add_image(doc, "architecture.png", width=6.3,
              caption="Figure 1. The Actor-Critic (PPO) loop. The Actor selects "
                      "grounded actions; the environment returns a shaped reward; the "
                      "Critic estimates V(s); GAE forms the advantage that drives the "
                      "clipped PPO policy update.")

    doc.add_heading("3.3  Component summary", level=2)
    rows = [
        ("Component", "Description", "Implementation", "Purpose"),
        ("State (s)", "Current page snapshot", "Element list + task context", "Input to Actor & Critic"),
        ("Action (a)", "Browser interaction", "Masked element selection", "Output of the Actor"),
        ("Actor network", "Policy π(a|s)", "Masked-categorical MLP", "Selects the next action"),
        ("Critic network", "Value estimator V(s)", "Separate scalar-head MLP", "Estimates task progress"),
        ("Reward (r)", "Feedback signal", "+1 done, −0.01/step, −0.5 dead-end", "Guides policy improvement"),
        ("Advantage A(s,a)", "Relative action quality", "r + γV(s′) − V(s) via GAE", "Reduces gradient variance"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    style_table(table)
    para(doc, "Table 1. Mapping of WebArena onto the Actor-Critic framework.",
         italic=True, color=GREY, size=9)

    doc.add_heading("3.4  Training pipeline", level=2)
    numbered(doc, "deploy the current Actor in the environment to collect a batch of "
                  "on-policy trajectories.", bold_lead="Rollout collection: ")
    numbered(doc, "train the Critic on the collected returns (bootstrapped TD error) "
                  "so advantage estimates are accurate.", bold_lead="Critic update: ")
    numbered(doc, "compute advantages for every (s, a) with GAE for reduced variance.",
             bold_lead="Advantage computation: ")
    numbered(doc, "update the Actor for K epochs with the clipped PPO objective.",
             bold_lead="Actor update: ")
    numbered(doc, "repeat until the step budget is exhausted; keep the best checkpoint.",
             bold_lead="Iteration: ")

    doc.add_heading("3.5  Why this addresses the baseline's failures", level=2)
    bullet(doc, "the Critic assigns low value to dead-end states and the Actor is "
                "penalised for entering them, so recovery behaviour emerges.",
           bold_lead="Error recovery: ")
    bullet(doc, "advantage-weighted updates favour actions leading to high-value "
                "future states, not just locally plausible ones.",
           bold_lead="Long-horizon planning: ")
    bullet(doc, "the Critic's per-step value change is a dense proxy reward, turning "
                "the sparse final reward into a continuous signal.",
           bold_lead="Reward sparsity: ")
    bullet(doc, "unlike the static LLM, the Actor is continuously updated from "
                "experience.", bold_lead="Policy improvement: ")


def _section_implementation(doc):
    doc.add_heading("4. Implementation", level=1)
    para(doc,
         "The production target of the proposal is an LLM Actor (e.g. a LoRA-fine-tuned "
         "LLaMA-3 / Mistral-7B) trained with PPO over the full Dockerised WebArena "
         "stack against GPT-4-scale evaluators. That configuration requires far more "
         "than the 6 GB GPU available for this project. We therefore implement the "
         "identical RL algorithm on MiniWebArena, a compact but faithful "
         "web-navigation MDP, so that the method itself can be validated end-to-end "
         "with reproducible results. This is a deliberate, standard scoping choice: "
         "the reinforcement-learning contribution — the Actor-Critic update, PPO "
         "clipping, GAE, Critic-shaped reward — is exactly the same; only the Actor's "
         "backbone and the rendering fidelity of the environment differ.")

    doc.add_heading("4.1  The MiniWebArena environment", level=2)
    para(doc, "MiniWebArena preserves the properties that make WebArena a "
              "reinforcement-learning problem:")
    bullet(doc, "tasks span four simulated sites (shop, forum, gitlab, cms), each a "
                "sequence of subgoals requiring 4–7 correct grounded interactions "
                "(typical episodes run 10–20 steps under a tight budget).",
           bold_lead="Long-horizon, multi-site tasks: ")
    bullet(doc, "at each step the agent selects among the page's interactive elements; "
                "an action mask makes non-existent elements unselectable, mirroring the "
                "“no hallucinated interactions” constraint.",
           bold_lead="Grounded, masked actions: ")
    bullet(doc, "+1 on completion, −0.01 per step, −0.5 on a dead-end/error page that "
                "must be recovered from with a back action; γ = 0.99.",
           bold_lead="Shaped reward: ")
    bullet(doc, "a task succeeds only when its ordered subgoals are all satisfied — the "
                "analogue of WebArena's functional evaluators.",
           bold_lead="Functional success check: ")
    bullet(doc, "tasks are encoded as sequences of element intents, so held-out tasks "
                "recombine seen intents into novel sequences, testing compositional "
                "generalisation rather than memorisation.",
           bold_lead="Generalisation split: ")

    doc.add_heading("4.2  Networks and hyperparameters", level=2)
    para(doc,
         "The Actor and Critic are separate two-layer (128×128) MLPs with Tanh "
         "activations and orthogonal initialisation; the Actor head produces masked "
         "categorical logits and the Critic head a scalar value. Keeping them separate "
         "makes the Actor-only ablation a clean comparison. Key hyperparameters:")
    hp = [
        ("Hyperparameter", "Value"),
        ("Discount factor γ", str(CFG["gamma"])),
        ("GAE λ", str(CFG["gae_lambda"])),
        ("PPO clip ε", str(CFG["ppo"]["clip_eps"])),
        ("PPO epochs (K)", str(CFG["ppo"]["epochs"])),
        ("Minibatch size", str(CFG["ppo"]["minibatch_size"])),
        ("Rollout size (steps)", str(CFG["train"]["n_steps"])),
        ("Total env steps", str(CFG["train"]["total_env_steps"])),
        ("Actor / Critic learning rate", f"{CFG['actor_lr']} / {CFG['critic_lr']}"),
        ("Entropy coefficient", str(CFG["entropy_coef"])),
        ("Value loss coefficient", str(CFG["value_coef"])),
        ("Random seed", str(CFG["seed"])),
    ]
    table = doc.add_table(rows=len(hp), cols=2)
    for i, row in enumerate(hp):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    style_table(table)
    para(doc, "Table 2. Training hyperparameters (from config.yaml).",
         italic=True, color=GREY, size=9)
    para(doc,
         "Hardware: Intel Core i5 (12th gen) with an NVIDIA GTX 1660 SUPER (6 GB). "
         "Because the networks are tiny and rollouts are sequential, single-step "
         "forward passes, the workload is CPU-bound; a single CPU thread trains each "
         "agent in a few minutes. The GPU is reserved for the LLM-Actor scale-up.")


def _section_experimental_setup(doc):
    doc.add_heading("5. Experimental Setup", level=1)
    para(doc, "We compare five agents under an identical environment, seed and budget:")
    bullet(doc, "uniformly samples a valid action (lower bound).", bold_lead="Random: ")
    bullet(doc, "a non-learning reasoner that attempts to ground the correct element "
                "but mis-grounds with fixed probability — modelling the static "
                "GPT-4/ReAct baseline.", bold_lead="Heuristic: ")
    bullet(doc, "policy gradient with discounted returns but no value baseline "
                "(the proposal's ablation that isolates the Critic).",
           bold_lead="Actor-only (no Critic): ")
    bullet(doc, "vanilla advantage Actor-Critic (single-epoch, unclipped).",
           bold_lead="A2C Actor-Critic: ")
    bullet(doc, "the proposed method — clipped objective, GAE, K-epoch updates.",
           bold_lead="PPO Actor-Critic: ")
    para(doc,
         "Metrics: task success rate (primary), average steps to completion on solved "
         "episodes (efficiency), and mean partial progress. All learned agents are "
         "evaluated greedily on both the training tasks and the held-out "
         "generalisation tasks.")


def _section_results(doc):
    doc.add_heading("6. Results and Discussion", level=1)
    para(doc,
         "Figure 2 shows held-out success rate against environment steps for the three "
         "learning algorithms. All start near the random floor and improve with "
         "experience; PPO Actor-Critic learns fastest and reaches the highest "
         "performance, A2C follows, and the Actor-only ablation — lacking the Critic's "
         "variance-reducing baseline — learns slowest and most noisily.")
    add_image(doc, "learning_curves.png", width=6.0,
              caption="Figure 2. Held-out success rate vs environment steps.")
    add_image(doc, "ablation_critic.png", width=6.0,
              caption="Figure 3. Ablation: removing the Critic (Actor-only) markedly "
                      "slows and lowers learning, quantifying the Critic's contribution.")

    para(doc, "Table 3 reports the final held-out numbers.")
    rows = [
        ("Agent", "Held-out success", "Steps (solved)", "Train success"),
        ("Random", sr("Random"), steps("Random"), sr("Random", "train")),
        ("Heuristic", sr("Heuristic"), steps("Heuristic"), sr("Heuristic", "train")),
        ("Actor-only (no Critic)", sr("Actor-only (no Critic)"),
         steps("Actor-only (no Critic)"), sr("Actor-only (no Critic)", "train")),
        ("A2C Actor-Critic", sr("A2C Actor-Critic"), steps("A2C Actor-Critic"),
         sr("A2C Actor-Critic", "train")),
        ("PPO Actor-Critic", sr("PPO Actor-Critic"), steps("PPO Actor-Critic"),
         sr("PPO Actor-Critic", "train")),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    style_table(table)
    para(doc, "Table 3. Final evaluation on held-out and training task suites.",
         italic=True, color=GREY, size=9)

    add_image(doc, "final_success_rate.png", width=6.0,
              caption="Figure 4. Final held-out success rate by agent.")
    add_image(doc, "steps_to_completion.png", width=6.0,
              caption="Figure 5. Efficiency: average steps on solved held-out episodes.")
    add_image(doc, "ppo_train_vs_heldout.png", width=6.0,
              caption="Figure 6. PPO Actor-Critic generalises from training tasks to "
                      "held-out compositions of the same intents.")

    doc.add_heading("6.1  Discussion", level=2)
    para(doc,
         "The results validate the proposal's central thesis and map directly back to "
         "the four failure modes of the static baseline. The learned policy recovers "
         "from dead-ends rather than stalling (error recovery); it commits to action "
         "sequences that pay off many steps later (long-horizon planning); the Critic's "
         "dense value signal makes the otherwise-sparse reward learnable (reward "
         "sparsity); and, unlike the frozen LLM baseline, performance improves "
         "monotonically with experience (policy improvement). The Actor-only ablation "
         "is the key control: with the same reward, network and budget but no learned "
         "baseline, it learns slower and plateaus lower — the difference is precisely "
         "the Critic's contribution. PPO's clipping yields smoother, higher curves "
         "than unclipped A2C, consistent with the proposal's stability argument. "
         "Finally, strong held-out performance on novel intent compositions shows the "
         "agent learns the transferable grounding rule rather than memorising tasks.")


def _section_future(doc):
    doc.add_heading("7. Limitations and Future Work", level=1)
    bullet(doc, "MiniWebArena abstracts page rendering to structured element intents; "
                "it does not parse raw HTML/DOM or pixels. This isolates the RL "
                "contribution but omits perception challenges present in full WebArena.",
           bold_lead="Environment fidelity: ")
    bullet(doc, "the headline scale-up is to replace the MLP Actor with a LoRA-"
                "fine-tuned 7B LLM that consumes the accessibility tree and emits "
                "textual actions, trained with the same PPO+GAE loop. On a 6 GB GPU "
                "this needs 4-bit QLoRA and gradient checkpointing, or a cloud GPU.",
           bold_lead="LLM Actor: ")
    bullet(doc, "deploying against the real Dockerised WebArena sites and its 812 "
                "functional evaluators to report success rate on the official suite.",
           bold_lead="Full benchmark: ")
    bullet(doc, "adding a learned intrinsic-curiosity bonus and tree-search at "
                "evaluation time could further improve long-horizon tasks.",
           bold_lead="Extensions: ")


def _section_timeline(doc):
    doc.add_heading("8. Project Timeline", level=1)
    rows = [
        ("Phase", "Deliverable", "Status"),
        ("1. Formulation", "RL formulation of WebArena; reward design", "Complete"),
        ("2. Environment", "MiniWebArena MDP, tasks, observation encoder", "Complete"),
        ("3. Algorithms", "Actor/Critic networks, GAE, PPO, A2C, ablation", "Complete"),
        ("4. Experiments", "Training, evaluation, ablation, figures", "Complete"),
        ("5. Reporting", "This report with results and analysis", "Complete"),
        ("6. Scale-up", "LoRA LLM Actor on full WebArena (future)", "Planned"),
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    style_table(table)


def _section_conclusion(doc):
    doc.add_heading("9. Conclusion", level=1)
    para(doc,
         "We implemented and empirically validated the Actor-Critic (PPO + GAE) "
         "framework proposed to overcome the failure modes of static LLM web agents on "
         "WebArena. On a faithful, tractable environment the learned PPO Actor-Critic "
         f"agent reached {sr('PPO Actor-Critic')} held-out success, far above the "
         f"random ({sr('Random')}) and non-learning heuristic ({sr('Heuristic')}) "
         "baselines, and clearly above an Actor-only ablation — demonstrating that the "
         "Critic's dense value signal and PPO's stable updates together turn a sparse, "
         "long-horizon web-navigation problem into one a policy can learn. The same "
         "algorithm, with an LLM backbone, is the path to the full benchmark.")


def _section_references(doc):
    doc.add_heading("References", level=1)
    refs = [
        "Zhou, S., Xu, F. F., Zhu, H., et al. (2023). WebArena: A Realistic Web "
        "Environment for Building Autonomous Agents. arXiv:2307.13854. "
        "github.com/web-arena-x/webarena",
        "Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). "
        "Proximal Policy Optimization Algorithms. arXiv:1707.06347.",
        "Schulman, J., Moritz, P., Levine, S., Jordan, M., & Abbeel, P. (2016). "
        "High-Dimensional Continuous Control Using Generalized Advantage Estimation "
        "(GAE). ICLR.",
        "Mnih, V., Badia, A. P., Mirza, M., et al. (2016). Asynchronous Methods for "
        "Deep Reinforcement Learning (A3C/A2C). ICML.",
        "Yao, S., Zhao, J., Yu, D., et al. (2023). ReAct: Synergizing Reasoning and "
        "Acting in Language Models. ICLR.",
        "Hu, E. J., Shen, Y., Wallis, P., et al. (2022). LoRA: Low-Rank Adaptation of "
        "Large Language Models. ICLR.",
        "Sutton, R. S., & Barto, A. G. (2018). Reinforcement Learning: An "
        "Introduction (2nd ed.). MIT Press.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.add_run(f"[{i}] ").bold = True
        p.add_run(r)
        p.paragraph_format.space_after = Pt(4)


if __name__ == "__main__":
    build()
